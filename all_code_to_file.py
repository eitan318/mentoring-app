import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pygments import lex
from pygments.lexers import get_lexer_for_filename, ClassNotFound
from pygments.lexers.dotnet import CSharpLexer
from pygments.lexers.html import XmlLexer
from pygments.lexers.templates import HtmlPhpLexer # Fallback/base template style or custom
from pygments.styles import get_style_by_name

# ==========================================
# CONFIGURATION
# ==========================================
SOURCE_DIR = r"C:\Users\eitan\source\repos\mentoring-app"
OUTPUT_FILE = r"C:\Users\eitan\source\repos\mentoring-app\all_code.docx"

# Ignore patterns
IGNORED_DIRS = {"obj", "bin", ".git", "node_modules", ".vs", ".claude", "Tests"}
IGNORED_FILES = {"AssemblyInfo.cs"}

# Allowed extensions
ALLOWED_EXTENSIONS = {".cs", ".xaml", ".razor"}

# Color palette theme for Pygments tokens (e.g., 'vs', 'default', 'monokai')
THEME_NAME = "vs" 


def set_cell_background(cell, color_hex):
    """Utility to give code blocks a light gray shading container."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def extract_class_name(file_path, extension):
    """Attempts to find a class, interface, or struct name within a file."""
    if extension != ".cs":
        return None
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            match = re.search(r"\b(class|interface|record|struct)\s+([a-zA-Z0-9_<>]+)", content)
            if match:
                return match.group(2)
    except Exception:
        pass
    return None


def add_code_block(doc, file_path, extension):
    """Appends syntax-highlighted code inside a shaded table cell (code block style).
    
    Fallback chain:
      1. Pygments syntax highlighting (full color + bold/italic)
      2. Plain monospace text (if Pygments fails or file unreadable)
    """
    def _apply_highlighted(p, code, file_path, extension):
        """Try to apply Pygments syntax highlighting. Returns True on success."""
        try:
            try:
                lexer = get_lexer_for_filename(file_path)
            except ClassNotFound:
                if extension in {".xaml", ".razor"}:
                    from pygments.lexers.html import XmlLexer
                    lexer = XmlLexer()
                else:
                    from pygments.lexers.special import TextLexer
                    lexer = TextLexer()

            style = get_style_by_name(THEME_NAME)
            tokens = list(lex(code, lexer))   # materialise so any lex error fires here

            for ttype, value in tokens:
                token_style = style.style_for_token(ttype)
                run = p.add_run(value)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)

                if token_style.get('color'):
                    hex_color = token_style['color']
                    try:
                        run.font.color.rgb = RGBColor(
                            int(hex_color[0:2], 16),
                            int(hex_color[2:4], 16),
                            int(hex_color[4:6], 16)
                        )
                    except (ValueError, IndexError):
                        pass  # Malformed colour — leave default

                if token_style.get('bold'):
                    run.bold = True
                if token_style.get('italic'):
                    run.italic = True

            return True

        except Exception:
            return False

    def _apply_plain(p, code):
        """Fallback: dump the whole file as unstyled monospace text."""
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    # ── Read file ────────────────────────────────────────────────────────────
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            code = f.read()
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f"[Error loading file '{file_path}': {e}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        return

    # ── Build shaded table cell ───────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F4F5")

    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    # ── Highlight → fallback ──────────────────────────────────────────────────
    if not _apply_highlighted(p, code, file_path, extension):
        _apply_plain(p, code)


def should_ignore(path: Path) -> bool:
    for part in path.parts:
        if part in IGNORED_DIRS:
            return True
    if path.name in IGNORED_FILES or path.name.endswith(".designer.cs"):
        return True
    return False


def main():
    doc = Document()
    
    # Base setup
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Simple title marker
    doc.add_heading("Project Code Architecture Export", level=0)
    p = doc.add_paragraph("Target Directory: ")
    p.add_run(SOURCE_DIR).bold = True
    doc.add_page_break()

    source_path = Path(SOURCE_DIR)
    file_count = 0

    print("Beginning folder traversal...")

    for root, dirs, files in os.walk(SOURCE_DIR):
        dirs[:] = [d for d in dirs if d not in IGNORED_DIRS]

        for file in files:
            file_path = Path(root) / file
            
            if file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
                continue
                
            if should_ignore(file_path):
                continue

            file_count += 1
            relative_path = file_path.relative_to(source_path)
            print(f"Adding: {relative_path}")

            # 1. File Path Header
            doc.add_heading(f"📄 File: {relative_path}", level=2)
            
            # 2. Add Class Banner if it exists
            class_name = extract_class_name(file_path, file_path.suffix.lower())
            if class_name:
                cp = doc.add_paragraph()
                c_run = cp.add_run(f"Type Core Identification: {class_name}")
                c_run.bold = True
                c_run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
                cp.paragraph_format.space_after = Pt(4)

            # 3. Inject Syntax-highlighted code cell block
            add_code_block(doc, file_path, file_path.suffix.lower())
            
            # Pad and segment to a clean next page
            doc.add_paragraph()
            doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f"\nCompleted! Generated clean copy of {file_count} files into:\n{OUTPUT_FILE}")


if __name__ == "__main__":
    main()